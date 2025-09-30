#!/usr/bin/env python
"""
创建示例 Word 模板文件

该脚本会创建几个常用的 Word 模板，包含 Jinja2 语法标记，
可以与 docxtpl MCP 服务器一起使用。
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import os


def create_invoice_template():
    """创建发票模板"""
    doc = Document()

    # 添加标题
    title = doc.add_heading('INVOICE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 公司信息
    doc.add_paragraph('{{company_name}}')
    doc.add_paragraph('{{company_address}}')
    doc.add_paragraph('Email: {{company_email}} | Phone: {{company_phone}}')

    doc.add_paragraph()

    # 客户信息
    doc.add_heading('Bill To:', level=2)
    doc.add_paragraph('{{customer_name}}')
    doc.add_paragraph('{{customer_address}}')
    doc.add_paragraph('{{customer_email}}')

    # 发票信息
    p = doc.add_paragraph()
    p.add_run('Invoice Number: ').bold = True
    p.add_run('{{invoice_number}}')

    p = doc.add_paragraph()
    p.add_run('Invoice Date: ').bold = True
    p.add_run('{{invoice_date|date}}')

    p = doc.add_paragraph()
    p.add_run('Due Date: ').bold = True
    p.add_run('{{due_date|date}}')

    doc.add_paragraph()

    # 添加商品表格
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'

    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Quantity'
    hdr_cells[3].text = 'Unit Price'
    hdr_cells[4].text = 'Total'

    # 添加 Jinja2 循环来填充表格
    doc.add_paragraph('{% for item in items %}')
    row_cells = table.add_row().cells
    row_cells[0].text = '{{loop.index}}'
    row_cells[1].text = '{{item.description}}'
    row_cells[2].text = '{{item.quantity}}'
    row_cells[3].text = '{{item.unit_price|currency}}'
    row_cells[4].text = '{{item.total|currency}}'
    doc.add_paragraph('{% endfor %}')

    doc.add_paragraph()

    # 总计
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('Subtotal: ').bold = True
    p.add_run('{{subtotal|currency}}')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('Tax ({{tax_rate}}%): ').bold = True
    p.add_run('{{tax_amount|currency}}')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('Total: ').bold = True
    p.add_run('{{total|currency}}')

    # 备注
    doc.add_paragraph()
    doc.add_heading('Notes:', level=2)
    doc.add_paragraph('{{notes}}')

    # 条款
    doc.add_paragraph()
    doc.add_heading('Terms & Conditions:', level=2)
    doc.add_paragraph('{{terms}}')

    # 保存
    doc.save('templates/invoice.docx')
    print('✅ 创建发票模板: templates/invoice.docx')


def create_report_template():
    """创建报告模板"""
    doc = Document()

    # 封面
    title = doc.add_heading('{{report_title}}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('{{report_subtitle}}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # 作者信息
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Prepared by: ').bold = True
    p.add_run('{{author_name}}')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('{{department}}')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('{{report_date|date}}')

    # 分页符
    doc.add_page_break()

    # 执行摘要
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph('{{executive_summary}}')

    doc.add_page_break()

    # 目录
    doc.add_heading('Table of Contents', 1)
    doc.add_paragraph('{% for section in sections %}')
    doc.add_paragraph('{{loop.index}}. {{section.title}}', style='List Number')
    doc.add_paragraph('{% endfor %}')

    doc.add_page_break()

    # 主要内容
    doc.add_paragraph('{% for section in sections %}')
    doc.add_heading('{{loop.index}}. {{section.title}}', 1)
    doc.add_paragraph('{{section.content}}')

    # 如果有子节
    doc.add_paragraph('{% if section.subsections %}')
    doc.add_paragraph('{% for subsection in section.subsections %}')
    doc.add_heading('{{loop.index}}.{{subsection.number}} {{subsection.title}}', 2)
    doc.add_paragraph('{{subsection.content}}')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph('{% endif %}')

    # 如果有数据表
    doc.add_paragraph('{% if section.table %}')
    doc.add_heading('Table: {{section.table.title}}', 3)
    table = doc.add_table(rows=1, cols=len('{{section.table.headers}}'))
    table.style = 'Light Shading Accent 1'
    doc.add_paragraph('{% endif %}')

    doc.add_paragraph('{% endfor %}')

    # 结论
    doc.add_page_break()
    doc.add_heading('Conclusions', 1)
    doc.add_paragraph('{{conclusions}}')

    # 建议
    doc.add_heading('Recommendations', 1)
    doc.add_paragraph('{% for recommendation in recommendations %}')
    doc.add_paragraph('{{loop.index}}. {{recommendation}}', style='List Number')
    doc.add_paragraph('{% endfor %}')

    # 附录
    doc.add_page_break()
    doc.add_heading('Appendices', 1)
    doc.add_paragraph('{{appendices}}')

    # 保存
    doc.save('templates/report.docx')
    print('✅ 创建报告模板: templates/report.docx')


def create_contract_template():
    """创建合同模板"""
    doc = Document()

    # 标题
    title = doc.add_heading('{{contract_type}} CONTRACT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 合同编号和日期
    p = doc.add_paragraph()
    p.add_run('Contract No: ').bold = True
    p.add_run('{{contract_number}}')

    p = doc.add_paragraph()
    p.add_run('Date: ').bold = True
    p.add_run('{{contract_date|date}}')

    doc.add_paragraph()

    # 当事方
    doc.add_heading('BETWEEN:', level=2)

    doc.add_paragraph('{{party1_name}}')
    doc.add_paragraph('{{party1_address}}')
    doc.add_paragraph('(hereinafter referred to as "{{party1_short_name}}")')

    doc.add_paragraph('AND')

    doc.add_paragraph('{{party2_name}}')
    doc.add_paragraph('{{party2_address}}')
    doc.add_paragraph('(hereinafter referred to as "{{party2_short_name}}")')

    doc.add_paragraph()

    # 背景
    doc.add_heading('WHEREAS:', level=2)
    doc.add_paragraph('{{whereas_clause}}')

    # 协议内容
    doc.add_heading('NOW, THEREFORE, in consideration of the mutual covenants and agreements hereinafter set forth, the parties agree as follows:', level=2)

    doc.add_paragraph()

    # 条款
    doc.add_paragraph('{% for clause in clauses %}')
    doc.add_heading('{{loop.index}}. {{clause.title}}', level=2)
    doc.add_paragraph('{{clause.content}}')

    # 子条款
    doc.add_paragraph('{% if clause.subclauses %}')
    doc.add_paragraph('{% for subclause in clause.subclauses %}')
    doc.add_paragraph('{{loop.index}}.{{subclause.number}} {{subclause.content}}', style='List Number')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph('{% endif %}')
    doc.add_paragraph('{% endfor %}')

    # 签名部分
    doc.add_page_break()
    doc.add_heading('IN WITNESS WHEREOF', level=2)
    doc.add_paragraph('The parties have executed this Agreement as of the date first above written.')

    doc.add_paragraph()
    doc.add_paragraph()

    # 签名块
    table = doc.add_table(rows=3, cols=2)

    # Party 1
    table.cell(0, 0).text = '{{party1_short_name}}:'
    table.cell(1, 0).text = '_' * 30
    table.cell(2, 0).text = 'Name: {{party1_signatory}}\nTitle: {{party1_title}}\nDate: {{signature_date1|date}}'

    # Party 2
    table.cell(0, 1).text = '{{party2_short_name}}:'
    table.cell(1, 1).text = '_' * 30
    table.cell(2, 1).text = 'Name: {{party2_signatory}}\nTitle: {{party2_title}}\nDate: {{signature_date2|date}}'

    # 保存
    doc.save('templates/contract.docx')
    print('✅ 创建合同模板: templates/contract.docx')


def create_letter_template():
    """创建信函模板"""
    doc = Document()

    # 发件人信息
    doc.add_paragraph('{{sender_name}}')
    doc.add_paragraph('{{sender_address}}')
    doc.add_paragraph('{{sender_city}}, {{sender_state}} {{sender_zip}}')
    doc.add_paragraph('{{sender_email}}')
    doc.add_paragraph('{{sender_phone}}')

    doc.add_paragraph()
    doc.add_paragraph('{{letter_date|date}}')
    doc.add_paragraph()

    # 收件人信息
    doc.add_paragraph('{{recipient_name}}')
    doc.add_paragraph('{{recipient_title}}')
    doc.add_paragraph('{{recipient_company}}')
    doc.add_paragraph('{{recipient_address}}')
    doc.add_paragraph('{{recipient_city}}, {{recipient_state}} {{recipient_zip}}')

    doc.add_paragraph()

    # 称呼
    doc.add_paragraph('Dear {{salutation}}:')

    doc.add_paragraph()

    # 主题（可选）
    doc.add_paragraph('{% if subject %}')
    p = doc.add_paragraph()
    p.add_run('RE: {{subject}}').bold = True
    doc.add_paragraph('{% endif %}')

    doc.add_paragraph()

    # 正文段落
    doc.add_paragraph('{% for paragraph in body_paragraphs %}')
    doc.add_paragraph('{{paragraph}}')
    doc.add_paragraph()
    doc.add_paragraph('{% endfor %}')

    # 结束语
    doc.add_paragraph('{{closing}},')

    doc.add_paragraph()
    doc.add_paragraph()

    # 签名
    doc.add_paragraph('{{sender_name}}')
    doc.add_paragraph('{{sender_title}}')

    # 附件（可选）
    doc.add_paragraph()
    doc.add_paragraph('{% if enclosures %}')
    doc.add_paragraph('Enclosures:')
    doc.add_paragraph('{% for enclosure in enclosures %}')
    doc.add_paragraph('- {{enclosure}}')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph('{% endif %}')

    # 抄送（可选）
    doc.add_paragraph('{% if cc_list %}')
    doc.add_paragraph('cc: {% for cc in cc_list %}{{cc}}{% if not loop.last %}, {% endif %}{% endfor %}')
    doc.add_paragraph('{% endif %}')

    # 保存
    doc.save('templates/letter.docx')
    print('✅ 创建信函模板: templates/letter.docx')


def create_templates_info():
    """创建模板使用说明文档"""
    info = """# Word 模板使用说明

## 可用模板列表

### 1. invoice.docx - 发票模板
**必需变量：**
- company_name: 公司名称
- company_address: 公司地址
- company_email: 公司邮箱
- company_phone: 公司电话
- customer_name: 客户名称
- customer_address: 客户地址
- customer_email: 客户邮箱
- invoice_number: 发票号
- invoice_date: 开票日期
- due_date: 到期日期
- items: 商品列表 (数组)
  - description: 商品描述
  - quantity: 数量
  - unit_price: 单价
  - total: 小计
- subtotal: 小计金额
- tax_rate: 税率
- tax_amount: 税额
- total: 总计
- notes: 备注
- terms: 条款

### 2. report.docx - 报告模板
**必需变量：**
- report_title: 报告标题
- report_subtitle: 副标题
- author_name: 作者姓名
- department: 部门
- report_date: 报告日期
- executive_summary: 执行摘要
- sections: 章节列表 (数组)
  - title: 章节标题
  - content: 章节内容
  - subsections: 子章节 (可选)
  - table: 数据表 (可选)
- conclusions: 结论
- recommendations: 建议列表 (数组)
- appendices: 附录

### 3. contract.docx - 合同模板
**必需变量：**
- contract_type: 合同类型
- contract_number: 合同编号
- contract_date: 合同日期
- party1_name: 甲方全称
- party1_address: 甲方地址
- party1_short_name: 甲方简称
- party2_name: 乙方全称
- party2_address: 乙方地址
- party2_short_name: 乙方简称
- whereas_clause: 鉴于条款
- clauses: 条款列表 (数组)
  - title: 条款标题
  - content: 条款内容
  - subclauses: 子条款 (可选)
- party1_signatory: 甲方签字人
- party1_title: 甲方职务
- party2_signatory: 乙方签字人
- party2_title: 乙方职务
- signature_date1: 甲方签字日期
- signature_date2: 乙方签字日期

### 4. letter.docx - 信函模板
**必需变量：**
- sender_name: 发件人姓名
- sender_address: 发件人地址
- sender_city: 发件人城市
- sender_state: 发件人州/省
- sender_zip: 发件人邮编
- sender_email: 发件人邮箱
- sender_phone: 发件人电话
- letter_date: 信函日期
- recipient_name: 收件人姓名
- recipient_title: 收件人职务
- recipient_company: 收件人公司
- recipient_address: 收件人地址
- recipient_city: 收件人城市
- recipient_state: 收件人州/省
- recipient_zip: 收件人邮编
- salutation: 称呼
- body_paragraphs: 正文段落 (数组)
- closing: 结束语
- sender_title: 发件人职务

**可选变量：**
- subject: 主题
- enclosures: 附件列表 (数组)
- cc_list: 抄送列表 (数组)

## 使用示例

### 生成发票
```json
{
  "template_name": "invoice.docx",
  "context_data": {
    "company_name": "ABC Tech Inc.",
    "company_address": "123 Main St, Suite 100",
    "company_email": "info@abctech.com",
    "company_phone": "+1-234-567-8900",
    "customer_name": "XYZ Corp",
    "customer_address": "456 Oak Ave",
    "customer_email": "billing@xyzcorp.com",
    "invoice_number": "INV-2024-001",
    "invoice_date": "2024-09-28",
    "due_date": "2024-10-28",
    "items": [
      {
        "description": "Software License",
        "quantity": 10,
        "unit_price": 99.99,
        "total": 999.90
      },
      {
        "description": "Support Service",
        "quantity": 1,
        "unit_price": 500.00,
        "total": 500.00
      }
    ],
    "subtotal": 1499.90,
    "tax_rate": 10,
    "tax_amount": 149.99,
    "total": 1649.89,
    "notes": "Thank you for your business!",
    "terms": "Payment due within 30 days"
  }
}
```

## 自定义过滤器

模板中可以使用以下自定义过滤器：

- **currency**: 格式化为货币显示（例：{{amount|currency}} → $1,234.56）
- **date**: 格式化日期（例：{{date|date}} → September 28, 2024）

## 注意事项

1. 所有模板文件必须放在 `templates` 目录下
2. 模板文件必须是 .docx 格式
3. 使用 {{variable}} 语法插入变量
4. 使用 {% for %} 语法进行循环
5. 使用 {% if %} 语法进行条件判断
6. 生成的文档会保存在 `output` 目录下
"""

    with open('templates/README.md', 'w', encoding='utf-8') as f:
        f.write(info)
    print('✅ 创建模板说明文档: templates/README.md')


def main():
    """主函数"""
    # 确保模板目录存在
    templates_dir = Path('templates')
    templates_dir.mkdir(exist_ok=True)

    print('🚀 开始创建 Word 模板...\n')

    try:
        # 创建各种模板
        create_invoice_template()
        create_report_template()
        create_contract_template()
        create_letter_template()

        # 创建说明文档
        create_templates_info()

        print('\n✨ 所有模板创建完成！')
        print('📁 模板位置: templates/')
        print('📖 查看 templates/README.md 了解如何使用这些模板')

    except Exception as e:
        print(f'❌ 创建模板时出错: {str(e)}')
        import traceback
        traceback.print_exc()


if __name__ == '__main__':
    main()