#!/usr/bin/env python
"""
docxtpl MCP Server

A Model Context Protocol server for generating Word documents from templates using docxtpl.
"""

import os
import sys
import json
import uuid
import base64
import logging
import traceback
from pathlib import Path
from datetime import datetime, date, timedelta
from typing import Dict, List, Any, Optional

from docxtpl import DocxTemplate
from docx import Document
from mcp.server import Server, NotificationOptions
from mcp.server.models import InitializationOptions
import mcp.types as types
import mcp.server.stdio

# Document parsing libraries
import pdfplumber
import docx2txt
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Get directories from environment or use defaults
TEMPLATE_DIR = Path(os.getenv('TEMPLATE_DIR', 'templates'))
OUTPUT_DIR = Path(os.getenv('OUTPUT_DIR', 'output'))
MAX_FILE_SIZE_MB = int(os.getenv('MAX_FILE_SIZE_MB', '50'))

# Ensure directories exist
TEMPLATE_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# Store generated documents metadata
generated_documents: Dict[str, Dict] = {}

class DocxTemplateServer:
    """Main MCP server for docxtpl operations"""

    def __init__(self):
        self.server = Server("docxtpl-mcp")
        self.setup_handlers()

    def setup_handlers(self):
        """Setup all MCP protocol handlers"""

        @self.server.list_tools()
        async def list_tools() -> List[types.Tool]:
            """List all available tools"""
            return [
                types.Tool(
                    name="generate_document",
                    description="Generate a Word document from a template with provided data",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "template_name": {
                                "type": "string",
                                "description": "Name of the template file (without path)"
                            },
                            "context_data": {
                                "type": "object",
                                "description": "JSON object containing the data to fill the template"
                            },
                            "output_name": {
                                "type": "string",
                                "description": "Optional output filename (without extension). If not provided, will use timestamp"
                            }
                        },
                        "required": ["template_name", "context_data"]
                    }
                ),
                types.Tool(
                    name="list_templates",
                    description="List all available Word templates",
                    inputSchema={
                        "type": "object",
                        "properties": {}
                    }
                ),
                types.Tool(
                    name="validate_template",
                    description="Validate a template and extract its variables",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "template_name": {
                                "type": "string",
                                "description": "Name of the template file to validate"
                            }
                        },
                        "required": ["template_name"]
                    }
                ),
                types.Tool(
                    name="preview_template",
                    description="Preview template with sample data to check rendering",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "template_name": {
                                "type": "string",
                                "description": "Name of the template file"
                            },
                            "sample_data": {
                                "type": "object",
                                "description": "Sample data to preview the template"
                            }
                        },
                        "required": ["template_name", "sample_data"]
                    }
                ),
                types.Tool(
                    name="delete_document",
                    description="Delete a generated document",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "document_id": {
                                "type": "string",
                                "description": "ID of the document to delete"
                            }
                        },
                        "required": ["document_id"]
                    }
                ),
                types.Tool(
                    name="list_documents",
                    description="List all generated documents",
                    inputSchema={
                        "type": "object",
                        "properties": {}
                    }
                ),
                types.Tool(
                    name="get_template_schema",
                    description="Get the complete schema for a template including all required and optional fields",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "template_name": {
                                "type": "string",
                                "description": "Name of the template file"
                            }
                        },
                        "required": ["template_name"]
                    }
                ),
                types.Tool(
                    name="generate_sample_data",
                    description="Generate sample data for a template with all required fields filled",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "template_name": {
                                "type": "string",
                                "description": "Name of the template file"
                            },
                            "locale": {
                                "type": "string",
                                "description": "Locale for sample data (en or zh). Default is en",
                                "enum": ["en", "zh"]
                            }
                        },
                        "required": ["template_name"]
                    }
                ),
                types.Tool(
                    name="parse_docx_document",
                    description="Parse a DOCX document and extract structured content including paragraphs, tables, and metadata",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Absolute path to the DOCX file to parse"
                            },
                            "include_tables": {
                                "type": "boolean",
                                "description": "Whether to extract tables from the document (default: true)"
                            }
                        },
                        "required": ["file_path"]
                    }
                ),
                types.Tool(
                    name="parse_pdf_document",
                    description="Parse a PDF document and extract text, tables, and metadata from each page",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Absolute path to the PDF file to parse"
                            },
                            "include_tables": {
                                "type": "boolean",
                                "description": "Whether to extract tables from the PDF (default: true)"
                            },
                            "pages": {
                                "type": "string",
                                "description": "Page range to parse (e.g., '1-5' or 'all'). Default is 'all'"
                            }
                        },
                        "required": ["file_path"]
                    }
                ),
                types.Tool(
                    name="extract_text_from_document",
                    description="Quick text extraction from DOCX, PDF, or Excel documents (without structure analysis)",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Absolute path to the document file (DOCX, PDF, or Excel)"
                            }
                        },
                        "required": ["file_path"]
                    }
                ),
                types.Tool(
                    name="get_document_metadata",
                    description="Extract metadata information from DOCX, PDF, or Excel documents",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Absolute path to the document file (DOCX, PDF, or Excel)"
                            }
                        },
                        "required": ["file_path"]
                    }
                ),
                types.Tool(
                    name="parse_excel_document",
                    description="Parse an Excel document (XLSX/XLS) and extract structured content including sheets, cells, and metadata",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Absolute path to the Excel file to parse"
                            },
                            "sheet_name": {
                                "type": "string",
                                "description": "Specific sheet name to parse (default: parse all sheets)"
                            },
                            "include_formulas": {
                                "type": "boolean",
                                "description": "Whether to include cell formulas (default: true)"
                            }
                        },
                        "required": ["file_path"]
                    }
                )
            ]

        @self.server.call_tool()
        async def call_tool(
            name: str,
            arguments: Dict[str, Any]
        ) -> List[types.TextContent | types.ImageContent | types.EmbeddedResource]:
            """Handle tool calls"""

            try:
                if name == "generate_document":
                    return await self.generate_document(
                        arguments.get("template_name"),
                        arguments.get("context_data"),
                        arguments.get("output_name")
                    )

                elif name == "list_templates":
                    return await self.list_templates()

                elif name == "validate_template":
                    return await self.validate_template(arguments.get("template_name"))

                elif name == "preview_template":
                    return await self.preview_template(
                        arguments.get("template_name"),
                        arguments.get("sample_data")
                    )

                elif name == "delete_document":
                    return await self.delete_document(arguments.get("document_id"))

                elif name == "list_documents":
                    return await self.list_documents()

                elif name == "get_template_schema":
                    return await self.get_template_schema(arguments.get("template_name"))

                elif name == "generate_sample_data":
                    return await self.generate_sample_data(
                        arguments.get("template_name"),
                        arguments.get("locale", "en")
                    )

                elif name == "parse_docx_document":
                    return await self.parse_docx_document(
                        arguments.get("file_path"),
                        arguments.get("include_tables", True)
                    )

                elif name == "parse_pdf_document":
                    return await self.parse_pdf_document(
                        arguments.get("file_path"),
                        arguments.get("include_tables", True),
                        arguments.get("pages", "all")
                    )

                elif name == "extract_text_from_document":
                    return await self.extract_text_from_document(
                        arguments.get("file_path")
                    )

                elif name == "get_document_metadata":
                    return await self.get_document_metadata(
                        arguments.get("file_path")
                    )

                elif name == "parse_excel_document":
                    return await self.parse_excel_document(
                        arguments.get("file_path"),
                        arguments.get("sheet_name"),
                        arguments.get("include_formulas", True)
                    )

                else:
                    return [types.TextContent(
                        type="text",
                        text=f"Unknown tool: {name}"
                    )]

            except Exception as e:
                logger.error(f"Error executing tool {name}: {str(e)}")
                return [types.TextContent(
                    type="text",
                    text=f"Error: {str(e)}\n\nTraceback:\n{traceback.format_exc()}"
                )]

        @self.server.list_resources()
        async def list_resources() -> List[types.Resource]:
            """List available resources"""
            resources = []

            # Add template resources
            for template_file in TEMPLATE_DIR.glob("*.docx"):
                resources.append(types.Resource(
                    uri=f"template://{template_file.stem}",
                    name=template_file.stem,
                    mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    description=f"Word template: {template_file.name}"
                ))

            # Add generated document resources
            for doc_id, doc_info in generated_documents.items():
                resources.append(types.Resource(
                    uri=f"document://{doc_id}",
                    name=doc_info["filename"],
                    mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    description=f"Generated document: {doc_info['filename']}"
                ))

            return resources

        @self.server.read_resource()
        async def read_resource(uri: str) -> str:
            """Read a specific resource"""

            if uri.startswith("template://"):
                template_name = uri.replace("template://", "")
                template_path = TEMPLATE_DIR / f"{template_name}.docx"

                if not template_path.exists():
                    return json.dumps({
                        "error": f"Template not found: {template_name}"
                    })

                # Return template information
                return json.dumps({
                    "name": template_name,
                    "path": str(template_path),
                    "size": template_path.stat().st_size,
                    "modified": template_path.stat().st_mtime
                })

            elif uri.startswith("document://"):
                doc_id = uri.replace("document://", "")

                if doc_id not in generated_documents:
                    return json.dumps({
                        "error": f"Document not found: {doc_id}"
                    })

                doc_info = generated_documents[doc_id]
                doc_path = Path(doc_info["path"])

                if not doc_path.exists():
                    return json.dumps({
                        "error": f"Document file not found: {doc_info['filename']}"
                    })

                # Return document information
                return json.dumps({
                    "id": doc_id,
                    "filename": doc_info["filename"],
                    "path": str(doc_path),
                    "size": doc_path.stat().st_size,
                    "created": doc_info["created"],
                    "template": doc_info["template"]
                })

            return json.dumps({
                "error": f"Unknown resource URI: {uri}"
            })

        @self.server.list_prompts()
        async def list_prompts() -> List[types.Prompt]:
            """List available prompts"""
            return [
                types.Prompt(
                    name="invoice_generator",
                    description="Generate an invoice from customer and product data",
                    arguments=[
                        types.PromptArgument(
                            name="customer_name",
                            description="Name of the customer",
                            required=True
                        ),
                        types.PromptArgument(
                            name="products",
                            description="Comma-separated list of products",
                            required=True
                        )
                    ]
                ),
                types.Prompt(
                    name="report_generator",
                    description="Generate a report with title, content and data",
                    arguments=[
                        types.PromptArgument(
                            name="title",
                            description="Report title",
                            required=True
                        ),
                        types.PromptArgument(
                            name="author",
                            description="Report author",
                            required=True
                        )
                    ]
                )
            ]

        @self.server.get_prompt()
        async def get_prompt(
            name: str,
            arguments: Optional[Dict[str, str]] = None
        ) -> types.GetPromptResult:
            """Get a specific prompt"""

            if name == "invoice_generator":
                customer = arguments.get("customer_name", "Customer")
                products = arguments.get("products", "Product 1, Product 2")

                return types.GetPromptResult(
                    description="Invoice generation prompt",
                    messages=[
                        types.PromptMessage(
                            role="user",
                            content=types.TextContent(
                                type="text",
                                text=f"""Please generate an invoice for {customer} with the following products: {products}.

Use the generate_document tool with the invoice template and provide appropriate context data including:
- Customer information (name, address, email)
- Invoice details (number, date, due date)
- Product list with quantities and prices
- Calculate totals and tax

Make the data realistic and professional."""
                            )
                        )
                    ]
                )

            elif name == "report_generator":
                title = arguments.get("title", "Report Title")
                author = arguments.get("author", "Author Name")

                return types.GetPromptResult(
                    description="Report generation prompt",
                    messages=[
                        types.PromptMessage(
                            role="user",
                            content=types.TextContent(
                                type="text",
                                text=f"""Please generate a report titled "{title}" by {author}.

Use the generate_document tool with the report template and provide appropriate context data including:
- Report metadata (title, author, date, version)
- Executive summary
- Main content sections
- Data tables or charts descriptions
- Conclusions and recommendations

Make the content professional and well-structured."""
                            )
                        )
                    ]
                )

            raise ValueError(f"Unknown prompt: {name}")

    async def generate_document(
        self,
        template_name: str,
        context_data: Dict[str, Any],
        output_name: Optional[str] = None
    ) -> List[types.TextContent]:
        """Generate a Word document from template"""

        # Validate template exists
        template_path = TEMPLATE_DIR / template_name
        if not template_path.exists():
            template_path = TEMPLATE_DIR / f"{template_name}.docx"
            if not template_path.exists():
                return [types.TextContent(
                    type="text",
                    text=f"Error: Template not found: {template_name}"
                )]

        # Generate output filename
        if output_name:
            output_filename = f"{output_name}.docx"
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{template_path.stem}_{timestamp}.docx"

        output_path = OUTPUT_DIR / output_filename

        try:
            # Load template and render with context
            doc = DocxTemplate(str(template_path))

            # Add some useful functions to the context
            context_data["now"] = datetime.now()
            context_data["today"] = datetime.now().date()

            # Custom filters
            def format_currency(value):
                try:
                    return f"${float(value):,.2f}"
                except:
                    return str(value)

            def format_date(value, format_str="%B %d, %Y"):
                if isinstance(value, str):
                    try:
                        value = datetime.fromisoformat(value)
                    except ValueError as e:
                        raise ValueError(f"Invalid date format for value '{value}'. Please use ISO format (YYYY-MM-DD)")
                elif not isinstance(value, (datetime, date)):
                    raise ValueError(f"Invalid date type: {type(value).__name__}. Expected string in ISO format or datetime object")
                return value.strftime(format_str)

            # Register custom filters using Jinja2 Environment
            from jinja2 import Environment
            jinja_env = Environment()
            jinja_env.filters['currency'] = format_currency
            jinja_env.filters['date'] = format_date

            # Render the document
            doc.render(context_data, jinja_env)

            # Save the document
            doc.save(str(output_path))

            # Check file size
            file_size_mb = output_path.stat().st_size / (1024 * 1024)
            if file_size_mb > MAX_FILE_SIZE_MB:
                output_path.unlink()
                return [types.TextContent(
                    type="text",
                    text=f"Error: Generated file exceeds maximum size ({MAX_FILE_SIZE_MB} MB)"
                )]

            # Generate document ID and store metadata
            doc_id = str(uuid.uuid4())[:8]
            generated_documents[doc_id] = {
                "id": doc_id,
                "filename": output_filename,
                "path": str(output_path),
                "template": template_name,
                "created": datetime.now().isoformat(),
                "size": output_path.stat().st_size
            }

            return [types.TextContent(
                type="text",
                text=f"""Document generated successfully!

📄 **File**: {output_filename}
📁 **Location**: {output_path}
🆔 **Document ID**: {doc_id}
📏 **Size**: {file_size_mb:.2f} MB
📋 **Template**: {template_name}
⏰ **Created**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

You can access this document using the resource URI: `document://{doc_id}`"""
            )]

        except Exception as e:
            logger.error(f"Error generating document: {str(e)}")
            error_message = str(e)

            # Clean template name for metadata lookup
            template_key = template_name.replace('.docx', '')

            # Try to provide more helpful error messages
            if "is undefined" in error_message:
                # Extract the undefined variable name
                import re
                match = re.search(r"'([^']+)' is undefined", error_message)
                if match:
                    missing_field = match.group(1)

                    # Load metadata to check if it's a required field
                    metadata_path = Path(__file__).parent.parent / "templates_metadata.json"
                    field_info = None
                    if metadata_path.exists():
                        with open(metadata_path, 'r', encoding='utf-8') as f:
                            metadata = json.load(f)
                            template_meta = metadata.get(template_key, {})
                            required_fields = template_meta.get("required_fields", {})
                            optional_fields = template_meta.get("optional_fields", {})

                            if missing_field in required_fields:
                                field_info = required_fields[missing_field]
                                field_type = "required"
                            elif missing_field in optional_fields:
                                field_info = optional_fields[missing_field]
                                field_type = "optional"

                    error_text = f"""❌ **Missing Field Error**

**Missing field**: `{missing_field}`
"""

                    if field_info:
                        error_text += f"""**Field type**: {field_type.capitalize()}
**Description**: {field_info.get('description', 'No description')}
**Data type**: {field_info.get('type', 'string')}
"""
                        if field_info.get('format'):
                            error_text += f"**Format**: {field_info['format']}\n"
                        if field_info.get('example'):
                            example = field_info['example']
                            if isinstance(example, list) or isinstance(example, dict):
                                example = json.dumps(example, ensure_ascii=False, indent=2)
                            error_text += f'\n**Example value**:\n```json\n"{missing_field}": {json.dumps(example) if not isinstance(example, str) else json.dumps(example)}\n```\n'
                    else:
                        error_text += f"\nThis field is required by the template but wasn't found in your data.\n"

                    error_text += "\n💡 **Tip**: Use `validate_template` to see all required fields."

                    return [types.TextContent(
                        type="text",
                        text=error_text
                    )]

            elif "Invalid isoformat string" in error_message:
                # Date format error
                import re
                match = re.search(r"Invalid isoformat string: '([^']+)'", error_message)
                invalid_date = match.group(1) if match else "unknown"

                return [types.TextContent(
                    type="text",
                    text=f"""❌ **Date Format Error**

**Invalid date value**: `{invalid_date}`

**Required format**: ISO date format (YYYY-MM-DD)
**Example**: "2025-01-28"

💡 **Common mistakes**:
- Using localized formats like "28/01/2025" or "Jan 28, 2025"
- Using Chinese date format like "2025年1月28日"

Please use the ISO format: YYYY-MM-DD"""
                )]

            elif "Object of type" in error_message and "is not JSON serializable" in error_message:
                # JSON serialization error - usually happens after another error
                return [types.TextContent(
                    type="text",
                    text=f"""❌ **Template Processing Error**

The template encountered an error during processing.
This often happens when:
1. A required field is missing
2. A date field has invalid format
3. An array field has incorrect structure

**Original error**: {error_message}

💡 **Debug tips**:
1. Run `validate_template` to see all required fields
2. Check that all date fields use ISO format (YYYY-MM-DD)
3. Ensure array fields contain proper object structures"""
                )]

            else:
                # Generic error with context
                # Don't include full context_data in error as it may contain datetime objects
                context_preview = {}
                for key, value in context_data.items():
                    if isinstance(value, (str, int, float, bool, type(None))):
                        context_preview[key] = value
                    elif isinstance(value, list) and len(value) > 0:
                        context_preview[key] = f"[Array with {len(value)} items]"
                    elif isinstance(value, dict):
                        context_preview[key] = f"[Object with {len(value)} keys]"
                    else:
                        context_preview[key] = f"[{type(value).__name__}]"

                return [types.TextContent(
                    type="text",
                    text=f"""❌ **Document Generation Error**

**Error**: {error_message}

**Template**: {template_name}

**Provided fields**:
{chr(10).join([f"- {k}: {v}" for k, v in context_preview.items()])}

💡 **Debug tips**:
1. Run `validate_template "{template_name}"` to see required fields
2. Use `generate_sample_data "{template_name}"` to get a working example
3. Check that all date fields use ISO format (YYYY-MM-DD)"""
                )]

    async def list_templates(self) -> List[types.TextContent]:
        """List all available templates"""

        templates = list(TEMPLATE_DIR.glob("*.docx"))

        if not templates:
            return [types.TextContent(
                type="text",
                text="No templates found. Please add .docx template files to the templates directory."
            )]

        template_info = []
        for template in templates:
            size_kb = template.stat().st_size / 1024
            modified = datetime.fromtimestamp(template.stat().st_mtime)
            template_info.append(
                f"- **{template.stem}** ({size_kb:.1f} KB) - Modified: {modified.strftime('%Y-%m-%d %H:%M')}"
            )

        return [types.TextContent(
            type="text",
            text=f"""📁 **Available Templates** ({len(templates)} found):

{chr(10).join(template_info)}

Use these template names with the `generate_document` tool."""
        )]

    async def validate_template(self, template_name: str) -> List[types.TextContent]:
        """Validate template and extract variables with detailed schema information"""

        # Clean template name (remove .docx extension if present)
        template_key = template_name.replace('.docx', '')

        # Check if template file exists
        template_path = TEMPLATE_DIR / f"{template_key}.docx"
        if not template_path.exists():
            return [types.TextContent(
                type="text",
                text=f"❌ Error: Template not found: {template_name}"
            )]

        try:
            # Load template metadata
            metadata_path = Path(__file__).parent.parent / "templates_metadata.json"
            if metadata_path.exists():
                with open(metadata_path, 'r', encoding='utf-8') as f:
                    metadata = json.load(f)
                    template_meta = metadata.get(template_key, {})
            else:
                template_meta = {}

            # If we have metadata, use it
            if template_meta:
                required_fields = template_meta.get("required_fields", {})
                optional_fields = template_meta.get("optional_fields", {})

                # Format required fields with type and example
                required_text = ""
                if required_fields:
                    required_text = "### 📌 Required Fields:\n"
                    for field_name, field_info in required_fields.items():
                        field_type = field_info.get("type", "string")
                        field_format = field_info.get("format", "")
                        description = field_info.get("description", "")
                        example = field_info.get("example", "")

                        format_hint = f" ({field_format})" if field_format else ""
                        required_text += f"- **{field_name}** (`{field_type}{format_hint}`): {description}\n"
                        if example:
                            if isinstance(example, list):
                                example_str = json.dumps(example, ensure_ascii=False, indent=2)
                                required_text += f"  Example: ```json\n{example_str}\n```\n"
                            else:
                                required_text += f"  Example: `{example}`\n"

                # Format optional fields
                optional_text = ""
                if optional_fields:
                    optional_text = "\n### 📋 Optional Fields:\n"
                    for field_name, field_info in optional_fields.items():
                        field_type = field_info.get("type", "string")
                        description = field_info.get("description", "")
                        optional_text += f"- **{field_name}** (`{field_type}`): {description}\n"

                # Count total fields
                total_required = len(required_fields)
                total_optional = len(optional_fields)

                return [types.TextContent(
                    type="text",
                    text=f"""✅ **Template Validation Results**:

📄 **Template**: {template_key}.docx
📏 **File Size**: {template_path.stat().st_size / 1024:.1f} KB
📝 **Description**: {template_meta.get("description", "No description available")}

**Field Summary**:
- Required Fields: {total_required}
- Optional Fields: {total_optional}
- Total Fields: {total_required + total_optional}

{required_text}
{optional_text}

### 📅 Date Format Note:
All date fields must be in ISO format: **YYYY-MM-DD** (e.g., "2025-01-28")

### 💡 Quick Start:
Use the `generate_sample_data` tool to get a complete example with all required fields filled."""
                )]

            # Fallback to old behavior if no metadata
            else:
                doc = DocxTemplate(str(template_path))
                variables = set()

                # Try to find variables by examining the XML
                try:
                    doc_xml = doc.get_xml()
                    if doc_xml:
                        import re
                        # Find Jinja2 variables
                        var_pattern = r'\{\{\s*(\w+)(?:\.\w+)*\s*\}\}'
                        found_vars = re.findall(var_pattern, doc_xml)
                        variables.update(found_vars)

                        # Find Jinja2 for loops
                        for_pattern = r'\{%\s*for\s+\w+\s+in\s+(\w+)\s*%\}'
                        found_loops = re.findall(for_pattern, doc_xml)
                        variables.update(found_loops)
                except Exception as scan_error:
                    logger.warning(f"Could not scan template for variables: {scan_error}")
                    variables.add("(Unable to scan - template may be corrupted)")

                return [types.TextContent(
                    type="text",
                    text=f"""⚠️ **Template Validation Results** (Basic Mode):

📄 **Template**: {template_key}.docx
📏 **File Size**: {template_path.stat().st_size / 1024:.1f} KB
🔤 **Variables Found**: {len(variables)}

**Template Variables**:
{chr(10).join([f"- {var}" for var in sorted(variables)])}

**Note**: Template metadata not found. Showing basic variable scan only.
For detailed field information, check templates_metadata.json"""
                )]

        except Exception as e:
            return [types.TextContent(
                type="text",
                text=f"❌ Error validating template: {str(e)}"
            )]

    async def preview_template(
        self,
        template_name: str,
        sample_data: Dict[str, Any]
    ) -> List[types.TextContent]:
        """Preview template with sample data"""

        # Generate a temporary preview
        preview_name = f"preview_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        result = await self.generate_document(template_name, sample_data, preview_name)

        # Add preview note to the result
        if "successfully" in result[0].text:
            result[0].text += "\n\n⚠️ **Note**: This is a preview document. It will be automatically cleaned up."

        return result

    async def delete_document(self, document_id: str) -> List[types.TextContent]:
        """Delete a generated document"""

        if document_id not in generated_documents:
            return [types.TextContent(
                type="text",
                text=f"Error: Document not found: {document_id}"
            )]

        doc_info = generated_documents[document_id]
        doc_path = Path(doc_info["path"])

        try:
            if doc_path.exists():
                doc_path.unlink()

            del generated_documents[document_id]

            return [types.TextContent(
                type="text",
                text=f"✅ Document deleted successfully: {doc_info['filename']}"
            )]

        except Exception as e:
            return [types.TextContent(
                type="text",
                text=f"Error deleting document: {str(e)}"
            )]

    async def list_documents(self) -> List[types.TextContent]:
        """List all generated documents"""

        if not generated_documents:
            return [types.TextContent(
                type="text",
                text="No documents generated yet. Use the `generate_document` tool to create documents."
            )]

        doc_list = []
        total_size = 0

        for doc_id, doc_info in generated_documents.items():
            doc_path = Path(doc_info["path"])
            if doc_path.exists():
                size_kb = doc_info["size"] / 1024
                total_size += doc_info["size"]
                created = datetime.fromisoformat(doc_info["created"])
                doc_list.append(
                    f"- **{doc_info['filename']}** (ID: `{doc_id}`)\n"
                    f"  - Template: {doc_info['template']}\n"
                    f"  - Size: {size_kb:.1f} KB\n"
                    f"  - Created: {created.strftime('%Y-%m-%d %H:%M:%S')}"
                )

        total_size_mb = total_size / (1024 * 1024)

        return [types.TextContent(
            type="text",
            text=f"""📚 **Generated Documents** ({len(generated_documents)} documents, {total_size_mb:.2f} MB total):

{chr(10).join(doc_list)}

Use document IDs with the `delete_document` tool to remove documents."""
        )]

    async def get_template_schema(self, template_name: str) -> List[types.TextContent]:
        """Get the complete schema for a template"""

        # Clean template name
        template_key = template_name.replace('.docx', '')

        # Check if template exists
        template_path = TEMPLATE_DIR / f"{template_key}.docx"
        if not template_path.exists():
            return [types.TextContent(
                type="text",
                text=f"❌ Error: Template not found: {template_name}"
            )]

        try:
            # Load metadata
            metadata_path = Path(__file__).parent.parent / "templates_metadata.json"
            if not metadata_path.exists():
                return [types.TextContent(
                    type="text",
                    text="❌ Error: Template metadata file not found. Run setup to create it."
                )]

            with open(metadata_path, 'r', encoding='utf-8') as f:
                metadata = json.load(f)

            if template_key not in metadata:
                return [types.TextContent(
                    type="text",
                    text=f"⚠️ No schema available for template: {template_name}\n\n"
                         f"Available templates: {', '.join(metadata.keys())}"
                )]

            schema = metadata[template_key]

            # Format as JSON schema
            json_schema = {
                "type": "object",
                "description": schema.get("description", ""),
                "properties": {},
                "required": []
            }

            # Add required fields
            for field_name, field_info in schema.get("required_fields", {}).items():
                json_schema["properties"][field_name] = {
                    "type": field_info.get("type", "string"),
                    "description": field_info.get("description", ""),
                    "example": field_info.get("example", "")
                }
                if field_info.get("format"):
                    json_schema["properties"][field_name]["format"] = field_info["format"]
                json_schema["required"].append(field_name)

            # Add optional fields
            for field_name, field_info in schema.get("optional_fields", {}).items():
                json_schema["properties"][field_name] = {
                    "type": field_info.get("type", "string"),
                    "description": field_info.get("description", ""),
                    "example": field_info.get("example", "")
                }
                if field_info.get("format"):
                    json_schema["properties"][field_name]["format"] = field_info["format"]

            return [types.TextContent(
                type="text",
                text=f"""📋 **Template Schema: {template_key}**

**Description**: {schema.get("description", "No description")}
**Total Fields**: {len(json_schema["properties"])} ({len(json_schema["required"])} required, {len(json_schema["properties"]) - len(json_schema["required"])} optional)

**JSON Schema**:
```json
{json.dumps(json_schema, indent=2, ensure_ascii=False)}
```

💡 **Usage Tip**: Use this schema to validate your data before generating documents."""
            )]

        except Exception as e:
            return [types.TextContent(
                type="text",
                text=f"❌ Error getting template schema: {str(e)}"
            )]

    async def generate_sample_data(
        self,
        template_name: str,
        locale: str = "en"
    ) -> List[types.TextContent]:
        """Generate sample data for a template"""

        # Clean template name
        template_key = template_name.replace('.docx', '')

        # Check if template exists
        template_path = TEMPLATE_DIR / f"{template_key}.docx"
        if not template_path.exists():
            return [types.TextContent(
                type="text",
                text=f"❌ Error: Template not found: {template_name}"
            )]

        try:
            # Load metadata
            metadata_path = Path(__file__).parent.parent / "templates_metadata.json"
            if not metadata_path.exists():
                return [types.TextContent(
                    type="text",
                    text="❌ Error: Template metadata file not found."
                )]

            with open(metadata_path, 'r', encoding='utf-8') as f:
                metadata = json.load(f)

            if template_key not in metadata:
                return [types.TextContent(
                    type="text",
                    text=f"⚠️ No schema available for template: {template_name}"
                )]

            schema = metadata[template_key]
            sample_data = {}

            # Generate sample data based on locale
            if locale == "zh":
                # Chinese sample data
                sample_data = self._generate_chinese_sample_data(template_key, schema)
            else:
                # English sample data (default)
                sample_data = self._generate_english_sample_data(template_key, schema)

            return [types.TextContent(
                type="text",
                text=f"""✅ **Sample Data for {template_key}** ({locale.upper()}):

```json
{json.dumps(sample_data, indent=2, ensure_ascii=False)}
```

💡 **Usage**:
1. Copy this JSON data
2. Modify values as needed
3. Use with `generate_document` tool:
   ```
   generate_document "{template_name}" <paste-modified-json>
   ```"""
            )]

        except Exception as e:
            return [types.TextContent(
                type="text",
                text=f"❌ Error generating sample data: {str(e)}"
            )]

    def _generate_english_sample_data(self, template_key: str, schema: Dict) -> Dict:
        """Generate English sample data for a template"""

        sample_data = {}

        # Use examples from metadata
        for field_name, field_info in schema.get("required_fields", {}).items():
            if field_info.get("example"):
                sample_data[field_name] = field_info["example"]
            else:
                # Generate default based on type
                field_type = field_info.get("type", "string")
                if field_type == "number":
                    sample_data[field_name] = 1000
                elif field_type == "array":
                    sample_data[field_name] = []
                elif field_type == "object":
                    sample_data[field_name] = {}
                elif field_info.get("format") == "date":
                    sample_data[field_name] = datetime.now().strftime("%Y-%m-%d")
                elif field_info.get("format") == "email":
                    sample_data[field_name] = "example@example.com"
                else:
                    sample_data[field_name] = f"Sample {field_name.replace('_', ' ').title()}"

        # Add some optional fields with examples
        for field_name, field_info in schema.get("optional_fields", {}).items():
            if field_info.get("example"):
                sample_data[field_name] = field_info["example"]

        return sample_data

    def _generate_chinese_sample_data(self, template_key: str, schema: Dict) -> Dict:
        """Generate Chinese sample data for a template"""

        # Template-specific Chinese sample data
        chinese_samples = {
            "invoice": {
                "company_name": "科技创新有限公司",
                "company_address": "北京市海淀区中关村科技园A座1001",
                "company_email": "info@keji.com.cn",
                "company_phone": "010-88889999",
                "customer_name": "张三",
                "customer_address": "上海市浦东新区陆家嘴金融中心",
                "customer_email": "zhangsan@example.cn",
                "invoice_number": "FP-2025-001",
                "invoice_date": datetime.now().strftime("%Y-%m-%d"),
                "due_date": (datetime.now() + timedelta(days=30)).strftime("%Y-%m-%d"),
                "items": [
                    {"description": "软件开发服务", "quantity": 1, "unit_price": 50000, "total": 50000},
                    {"description": "技术咨询服务", "quantity": 10, "unit_price": 2000, "total": 20000}
                ],
                "subtotal": 70000,
                "tax_rate": 0.06,
                "tax_amount": 4200,
                "total": 74200,
                "notes": "感谢您的合作！",
                "terms": "收到发票后30天内付款"
            },
            "letter": {
                "sender_name": "王明",
                "sender_title": "市场部经理",
                "sender_address": "北京市朝阳区建国路88号",
                "sender_city": "北京",
                "sender_state": "北京市",
                "sender_zip": "100022",
                "sender_email": "wangming@company.cn",
                "sender_phone": "138-0000-1234",
                "letter_date": datetime.now().strftime("%Y-%m-%d"),
                "recipient_name": "李华",
                "recipient_title": "总经理",
                "recipient_company": "创新科技有限公司",
                "recipient_address": "上海市浦东新区世纪大道100号",
                "recipient_city": "上海",
                "recipient_state": "上海市",
                "recipient_zip": "200120",
                "subject": "关于商务合作的提案",
                "salutation": "尊敬的李总",
                "body_paragraphs": [
                    "很高兴有机会向您介绍我们的最新产品和服务。",
                    "我们公司专注于提供高质量的技术解决方案，已经为众多客户提供了优质的服务。",
                    "期待能够与贵公司建立长期的合作关系。"
                ],
                "closing": "此致敬礼",
                "enclosures": ["产品介绍手册", "合作方案书"],
                "cc_list": ["销售总监", "技术总监"]
            },
            "report": {
                "report_title": "年度业绩报告",
                "report_subtitle": "2025财政年度",
                "author_name": "陈晓",
                "department": "财务部",
                "report_date": datetime.now().strftime("%Y-%m-%d"),
                "executive_summary": "本报告总结了公司2025年度的财务表现和主要成就。",
                "sections": [
                    {
                        "title": "概述",
                        "content": "2025年公司业绩稳步增长，营收同比增长25%。"
                    },
                    {
                        "title": "财务分析",
                        "content": "详细的财务数据分析显示各部门均实现了预定目标。"
                    }
                ],
                "conclusions": "公司整体表现优异，达到了年初制定的各项目标。",
                "recommendations": ["继续加大研发投入", "拓展海外市场", "优化成本结构"]
            },
            "contract": {
                "contract_type": "服务",
                "contract_number": "HT-2025-001",
                "contract_date": datetime.now().strftime("%Y-%m-%d"),
                "party1_name": "甲方科技有限公司",
                "party1_address": "北京市海淀区中关村大街1号",
                "party1_short_name": "甲方",
                "party2_name": "乙方服务有限公司",
                "party2_address": "上海市浦东新区张江高科技园区",
                "party2_short_name": "乙方",
                "whereas_clause": "鉴于甲方需要技术服务，乙方具备相应的技术能力和资质",
                "clauses": [
                    {
                        "title": "服务内容",
                        "content": "乙方为甲方提供软件开发和技术支持服务。",
                        "subclauses": [
                            {"number": "1", "content": "软件开发服务"},
                            {"number": "2", "content": "技术支持服务"}
                        ]
                    },
                    {
                        "title": "合作期限",
                        "content": "本合同有效期为一年，自签署之日起生效。"
                    }
                ],
                "party1_signatory": "张总",
                "party1_title": "总经理",
                "party2_signatory": "李总",
                "party2_title": "总经理",
                "signature_date1": datetime.now().strftime("%Y-%m-%d"),
                "signature_date2": datetime.now().strftime("%Y-%m-%d")
            }
        }

        # Use template-specific sample if available
        if template_key in chinese_samples:
            return chinese_samples[template_key]

        # Otherwise generate generic Chinese sample
        return self._generate_english_sample_data(template_key, schema)

    async def parse_docx_document(
        self,
        file_path: str,
        include_tables: bool = True
    ) -> List[types.TextContent]:
        """Parse a DOCX document and extract structured content"""

        try:
            doc_path = Path(file_path)

            # Validate file exists
            if not doc_path.exists():
                return [types.TextContent(
                    type="text",
                    text=f"❌ 错误: 文件不存在: {file_path}"
                )]

            # Validate file extension
            if doc_path.suffix.lower() != '.docx':
                return [types.TextContent(
                    type="text",
                    text=f"❌ 错误: 文件格式不正确。期望 .docx 文件,实际: {doc_path.suffix}"
                )]

            # Check file size
            file_size_mb = doc_path.stat().st_size / (1024 * 1024)
            if file_size_mb > MAX_FILE_SIZE_MB:
                return [types.TextContent(
                    type="text",
                    text=f"❌ 错误: 文件大小超过限制 ({MAX_FILE_SIZE_MB} MB)"
                )]

            # Parse document
            doc = Document(str(doc_path))

            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "filename": doc_path.name,
                "file_size_mb": round(file_size_mb, 2),
                "author": core_props.author or "Unknown",
                "title": core_props.title or "",
                "subject": core_props.subject or "",
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "last_modified_by": core_props.last_modified_by or "",
            }

            # Extract paragraphs with styles
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    para_info = {
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal"
                    }
                    paragraphs.append(para_info)

            # Extract tables if requested
            tables_data = []
            if include_tables and doc.tables:
                for table_idx, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)

                    tables_data.append({
                        "table_number": table_idx + 1,
                        "rows": len(table.rows),
                        "columns": len(table.columns),
                        "data": table_data
                    })

            # Construct result
            result = {
                "metadata": metadata,
                "content": {
                    "paragraphs": paragraphs,
                    "paragraph_count": len(paragraphs),
                    "tables": tables_data,
                    "table_count": len(tables_data)
                }
            }

            return [types.TextContent(
                type="text",
                text=f"""✅ **DOCX 文档解析成功**

📄 **文件信息**:
- 文件名: {metadata['filename']}
- 大小: {metadata['file_size_mb']} MB
- 作者: {metadata['author']}
- 标题: {metadata['title'] or '(无)'}

📝 **内容统计**:
- 段落数: {result['content']['paragraph_count']}
- 表格数: {result['content']['table_count']}

📋 **解析结果 (JSON)**:
```json
{json.dumps(result, indent=2, ensure_ascii=False)}
```

💡 提示: 使用此结构化数据可以进行进一步分析或转换。"""
            )]

        except Exception as e:
            logger.error(f"解析 DOCX 文档时出错: {str(e)}")
            return [types.TextContent(
                type="text",
                text=f"❌ **解析失败**: {str(e)}\n\n{traceback.format_exc()}"
            )]

    async def parse_pdf_document(
        self,
        file_path: str,
        include_tables: bool = True,
        pages: str = "all"
    ) -> List[types.TextContent]:
        """Parse a PDF document and extract text and tables"""

        try:
            pdf_path = Path(file_path)

            # Validate file exists
            if not pdf_path.exists():
                return [types.TextContent(
                    type="text",
                    text=f"❌ 错误: 文件不存在: {file_path}"
                )]

            # Validate file extension
            if pdf_path.suffix.lower() != '.pdf':
                return [types.TextContent(
                    type="text",
                    text=f"❌ 错误: 文件格式不正确。期望 .pdf 文件,实际: {pdf_path.suffix}"
                )]

            # Check file size
            file_size_mb = pdf_path.stat().st_size / (1024 * 1024)
            if file_size_mb > MAX_FILE_SIZE_MB:
                return [types.TextContent(
                    type="text",
                    text=f"❌ 错误: 文件大小超过限制 ({MAX_FILE_SIZE_MB} MB)"
                )]

            # Open PDF
            with pdfplumber.open(str(pdf_path)) as pdf:
                # Extract metadata
                metadata = {
                    "filename": pdf_path.name,
                    "file_size_mb": round(file_size_mb, 2),
                    "pages": len(pdf.pages),
                    "metadata": pdf.metadata or {}
                }

                # Parse page range
                if pages == "all":
                    page_indices = range(len(pdf.pages))
                else:
                    # Parse range like "1-5" or "1,3,5"
                    try:
                        if '-' in pages:
                            start, end = pages.split('-')
                            page_indices = range(int(start) - 1, int(end))
                        elif ',' in pages:
                            page_indices = [int(p.strip()) - 1 for p in pages.split(',')]
                        else:
                            page_indices = [int(pages) - 1]
                    except ValueError:
                        return [types.TextContent(
                            type="text",
                            text=f"❌ 错误: 无效的页面范围: {pages}"
                        )]

                # Extract content from pages
                pages_data = []
                for idx in page_indices:
                    if idx >= len(pdf.pages):
                        continue

                    page = pdf.pages[idx]
                    page_info = {
                        "page_number": idx + 1,
                        "text": page.extract_text() or "",
                        "width": page.width,
                        "height": page.height
                    }

                    # Extract tables if requested
                    if include_tables:
                        tables = page.extract_tables()
                        if tables:
                            page_info["tables"] = [
                                {
                                    "table_number": t_idx + 1,
                                    "rows": len(table),
                                    "columns": len(table[0]) if table else 0,
                                    "data": table
                                }
                                for t_idx, table in enumerate(tables)
                            ]
                        else:
                            page_info["tables"] = []
                    else:
                        page_info["tables"] = []

                    pages_data.append(page_info)

            # Construct result
            result = {
                "metadata": metadata,
                "pages": pages_data,
                "total_pages_parsed": len(pages_data)
            }

            # Calculate statistics
            total_tables = sum(len(p.get("tables", [])) for p in pages_data)
            total_text_length = sum(len(p.get("text", "")) for p in pages_data)

            return [types.TextContent(
                type="text",
                text=f"""✅ **PDF 文档解析成功**

📄 **文件信息**:
- 文件名: {metadata['filename']}
- 大小: {metadata['file_size_mb']} MB
- 总页数: {metadata['pages']}
- 已解析: {len(pages_data)} 页

📝 **内容统计**:
- 文本长度: {total_text_length} 字符
- 表格数: {total_tables}

📋 **解析结果 (JSON)**:
```json
{json.dumps(result, indent=2, ensure_ascii=False)}
```

💡 提示: 可以使用 pages 参数指定解析特定页面 (例如: "1-5" 或 "1,3,5")"""
            )]

        except Exception as e:
            logger.error(f"解析 PDF 文档时出错: {str(e)}")
            return [types.TextContent(
                type="text",
                text=f"❌ **解析失败**: {str(e)}\n\n{traceback.format_exc()}"
            )]

    async def extract_text_from_document(
        self,
        file_path: str
    ) -> List[types.TextContent]:
        """Quick text extraction from DOCX or PDF documents"""

        try:
            doc_path = Path(file_path)

            # Validate file exists
            if not doc_path.exists():
                return [types.TextContent(
                    type="text",
                    text=f"❌ 错误: 文件不存在: {file_path}"
                )]

            file_ext = doc_path.suffix.lower()

            # Extract text based on file type
            if file_ext == '.docx':
                # Use docx2txt for quick text extraction
                text = docx2txt.process(str(doc_path))

            elif file_ext == '.pdf':
                # Use pdfplumber for quick text extraction
                text_parts = []
                with pdfplumber.open(str(doc_path)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                text = "\n\n".join(text_parts)

            elif file_ext in ['.xlsx', '.xls']:
                # Use openpyxl for quick text extraction from Excel
                wb = load_workbook(str(doc_path), data_only=True, read_only=True)
                text_parts = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    text_parts.append(f"=== {sheet_name} ===\n")
                    for row in ws.iter_rows(values_only=True):
                        row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                        if row_text.strip():
                            text_parts.append(row_text)
                text = "\n".join(text_parts)
                wb.close()

            else:
                return [types.TextContent(
                    type="text",
                    text=f"❌ 错误: 不支持的文件格式: {file_ext}\n仅支持 .docx, .pdf, .xlsx 和 .xls 文件"
                )]

            # Statistics
            char_count = len(text)
            word_count = len(text.split())
            line_count = len(text.split('\n'))

            return [types.TextContent(
                type="text",
                text=f"""✅ **文本提取成功**

📄 **文件**: {doc_path.name}

📊 **统计信息**:
- 字符数: {char_count:,}
- 单词数: {word_count:,}
- 行数: {line_count:,}

📝 **提取的文本**:
```
{text[:2000]}{'...' if len(text) > 2000 else ''}
```

💡 提示: 如需完整结构化解析,请使用 parse_docx_document, parse_pdf_document 或 parse_excel_document"""
            )]

        except Exception as e:
            logger.error(f"提取文本时出错: {str(e)}")
            return [types.TextContent(
                type="text",
                text=f"❌ **提取失败**: {str(e)}\n\n{traceback.format_exc()}"
            )]

    async def get_document_metadata(
        self,
        file_path: str
    ) -> List[types.TextContent]:
        """Extract metadata from DOCX or PDF documents"""

        try:
            doc_path = Path(file_path)

            # Validate file exists
            if not doc_path.exists():
                return [types.TextContent(
                    type="text",
                    text=f"❌ 错误: 文件不存在: {file_path}"
                )]

            file_ext = doc_path.suffix.lower()
            file_size_mb = doc_path.stat().st_size / (1024 * 1024)

            metadata = {
                "filename": doc_path.name,
                "file_path": str(doc_path),
                "file_size_mb": round(file_size_mb, 2),
                "file_type": file_ext,
            }

            # Extract metadata based on file type
            if file_ext == '.docx':
                doc = Document(str(doc_path))
                core_props = doc.core_properties

                metadata.update({
                    "author": core_props.author or "Unknown",
                    "title": core_props.title or "",
                    "subject": core_props.subject or "",
                    "keywords": core_props.keywords or "",
                    "created": core_props.created.isoformat() if core_props.created else None,
                    "modified": core_props.modified.isoformat() if core_props.modified else None,
                    "last_modified_by": core_props.last_modified_by or "",
                    "revision": core_props.revision,
                    "category": core_props.category or "",
                    "comments": core_props.comments or "",
                })

                # Document statistics
                metadata["statistics"] = {
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "sections": len(doc.sections)
                }

            elif file_ext == '.pdf':
                with pdfplumber.open(str(doc_path)) as pdf:
                    pdf_metadata = pdf.metadata or {}

                    metadata.update({
                        "pages": len(pdf.pages),
                        "author": pdf_metadata.get('Author', 'Unknown'),
                        "title": pdf_metadata.get('Title', ''),
                        "subject": pdf_metadata.get('Subject', ''),
                        "creator": pdf_metadata.get('Creator', ''),
                        "producer": pdf_metadata.get('Producer', ''),
                        "created": pdf_metadata.get('CreationDate', ''),
                        "modified": pdf_metadata.get('ModDate', ''),
                    })

            elif file_ext in ['.xlsx', '.xls']:
                wb = load_workbook(str(doc_path), data_only=True, read_only=True)

                metadata.update({
                    "sheets_count": len(wb.sheetnames),
                    "sheet_names": wb.sheetnames,
                })

                # Add workbook properties if available
                if wb.properties:
                    props = wb.properties
                    metadata.update({
                        "creator": props.creator or "Unknown",
                        "title": props.title or "",
                        "subject": props.subject or "",
                        "description": props.description or "",
                        "keywords": props.keywords or "",
                        "created": props.created.isoformat() if props.created else None,
                        "modified": props.modified.isoformat() if props.modified else None,
                        "last_modified_by": props.lastModifiedBy or "",
                        "category": props.category or "",
                    })

                # Calculate statistics
                total_cells = 0
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    if ws.max_row and ws.max_column:
                        total_cells += ws.max_row * ws.max_column

                metadata["statistics"] = {
                    "total_cells": total_cells
                }

                wb.close()

            else:
                return [types.TextContent(
                    type="text",
                    text=f"❌ 错误: 不支持的文件格式: {file_ext}\n仅支持 .docx, .pdf, .xlsx 和 .xls 文件"
                )]

            return [types.TextContent(
                type="text",
                text=f"""✅ **文档元数据**

📄 **基本信息**:
- 文件名: {metadata['filename']}
- 大小: {metadata['file_size_mb']} MB
- 类型: {metadata['file_type']}

📋 **完整元数据 (JSON)**:
```json
{json.dumps(metadata, indent=2, ensure_ascii=False)}
```"""
            )]

        except Exception as e:
            logger.error(f"获取元数据时出错: {str(e)}")
            return [types.TextContent(
                type="text",
                text=f"❌ **获取失败**: {str(e)}\n\n{traceback.format_exc()}"
            )]

    async def parse_excel_document(
        self,
        file_path: str,
        sheet_name: Optional[str] = None,
        include_formulas: bool = True
    ) -> List[types.TextContent]:
        """Parse an Excel document and extract structured content"""

        try:
            excel_path = Path(file_path)

            # Validate file exists
            if not excel_path.exists():
                return [types.TextContent(
                    type="text",
                    text=f"❌ 错误: 文件不存在: {file_path}"
                )]

            # Validate file extension
            if excel_path.suffix.lower() not in ['.xlsx', '.xls']:
                return [types.TextContent(
                    type="text",
                    text=f"❌ 错误: 文件格式不正确。期望 .xlsx 或 .xls 文件,实际: {excel_path.suffix}"
                )]

            # Check file size
            file_size_mb = excel_path.stat().st_size / (1024 * 1024)
            if file_size_mb > MAX_FILE_SIZE_MB:
                return [types.TextContent(
                    type="text",
                    text=f"❌ 错误: 文件大小超过限制 ({MAX_FILE_SIZE_MB} MB)"
                )]

            # Load workbook (read_only=False to access merged_cells)
            wb = load_workbook(str(excel_path), data_only=False, read_only=False)

            # Extract metadata
            metadata = {
                "filename": excel_path.name,
                "file_size_mb": round(file_size_mb, 2),
                "sheets_count": len(wb.sheetnames),
                "sheet_names": wb.sheetnames,
            }

            # Add workbook properties if available
            if wb.properties:
                props = wb.properties
                metadata.update({
                    "creator": props.creator or "Unknown",
                    "title": props.title or "",
                    "subject": props.subject or "",
                    "description": props.description or "",
                    "created": props.created.isoformat() if props.created else None,
                    "modified": props.modified.isoformat() if props.modified else None,
                })

            # Determine which sheets to parse
            if sheet_name:
                if sheet_name not in wb.sheetnames:
                    wb.close()
                    return [types.TextContent(
                        type="text",
                        text=f"❌ 错误: 工作表 '{sheet_name}' 不存在\n可用工作表: {', '.join(wb.sheetnames)}"
                    )]
                sheets_to_parse = [sheet_name]
            else:
                sheets_to_parse = wb.sheetnames

            # Parse sheets
            sheets_data = []
            for ws_name in sheets_to_parse:
                ws = wb[ws_name]

                # Get sheet dimensions
                if ws.max_row is None or ws.max_column is None:
                    continue

                # Extract cell data
                data = []
                for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=ws.max_column):
                    row_data = []
                    for cell in row:
                        cell_value = cell.value
                        # Convert datetime to ISO format string
                        if isinstance(cell_value, (datetime, date)):
                            cell_value = cell_value.isoformat()
                        row_data.append(cell_value)
                    data.append(row_data)

                # Extract formulas if requested
                formulas = {}
                if include_formulas:
                    for row in ws.iter_rows():
                        for cell in row:
                            if cell.value and isinstance(cell.value, str) and cell.value.startswith('='):
                                cell_ref = f"{get_column_letter(cell.column)}{cell.row}"
                                formulas[cell_ref] = cell.value

                # Extract merged cells
                merged_cells = []
                if ws.merged_cells:
                    merged_cells = [str(merged_range) for merged_range in ws.merged_cells.ranges]

                sheet_info = {
                    "name": ws_name,
                    "rows": ws.max_row,
                    "columns": ws.max_column,
                    "data": data,
                    "merged_cells": merged_cells,
                }

                if formulas:
                    sheet_info["formulas"] = formulas

                sheets_data.append(sheet_info)

            wb.close()

            # Construct result
            result = {
                "metadata": metadata,
                "sheets": sheets_data,
                "total_sheets_parsed": len(sheets_data)
            }

            # Calculate statistics
            total_cells = sum(s["rows"] * s["columns"] for s in sheets_data)
            total_formulas = sum(len(s.get("formulas", {})) for s in sheets_data)

            return [types.TextContent(
                type="text",
                text=f"""✅ **Excel 文档解析成功**

📄 **文件信息**:
- 文件名: {metadata['filename']}
- 大小: {metadata['file_size_mb']} MB
- 工作表总数: {metadata['sheets_count']}
- 已解析: {len(sheets_data)} 个工作表

📊 **内容统计**:
- 工作表名称: {', '.join([s['name'] for s in sheets_data])}
- 总单元格数: {total_cells:,}
- 公式数: {total_formulas}

📋 **解析结果 (JSON)**:
```json
{json.dumps(result, indent=2, ensure_ascii=False)}
```

💡 提示: 可以使用 sheet_name 参数指定解析特定工作表"""
            )]

        except Exception as e:
            logger.error(f"解析 Excel 文档时出错: {str(e)}")
            return [types.TextContent(
                type="text",
                text=f"❌ **解析失败**: {str(e)}\n\n{traceback.format_exc()}"
            )]

    async def run(self):
        """Run the MCP server"""
        async with mcp.server.stdio.stdio_server() as (read_stream, write_stream):
            init_options = InitializationOptions(
                server_name="docxtpl-mcp",
                server_version="0.1.0",
                capabilities=self.server.get_capabilities(
                    notification_options=NotificationOptions(),
                    experimental_capabilities={},
                )
            )

            logger.info("Starting docxtpl MCP server...")
            logger.info(f"Template directory: {TEMPLATE_DIR}")
            logger.info(f"Output directory: {OUTPUT_DIR}")

            await self.server.run(
                read_stream,
                write_stream,
                init_options
            )


async def main():
    """Main entry point"""
    server = DocxTemplateServer()
    await server.run()


if __name__ == "__main__":
    import asyncio
    asyncio.run(main())